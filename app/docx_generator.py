import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from schema.models import TailoredCVOutput


def _add_section_heading(doc: Document, text: str):
    """Utility to add consistent ATS-style section headings with bottom borders."""
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(4)
    heading.paragraph_format.keep_with_next = True

    run = heading.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 51, 102)  # Deep Navy

    # Add subtle horizontal line below section header
    pPr = heading._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")  # Border thickness
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "003366")
    pBdr.append(bottom)
    pPr.append(pBdr)


def create_docx(
    tailored_cv: TailoredCVOutput, 
    output_dir: str = "generated_cvs"
) -> str:
    """
    Converts a TailoredCVOutput object into a formatted ATS-compliant .docx file.
    Returns the file path of the generated document.
    """
    # Ensure destination directory exists
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()

    # Set standard 0.75-inch margins (Optimal for ATS readability)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Base styling
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(51, 51, 51)  # Off-black for crisp readability

    # -------------------------------------------------------------
    # 1. HEADER / CONTACT INFORMATION
    # -------------------------------------------------------------
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.paragraph_format.space_after = Pt(2)

    # Name
    name_run = header_p.add_run(tailored_cv.full_name)
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = RGBColor(0, 51, 102)

    # Contact Info Line
    contact_parts = []
    if tailored_cv.email:
        contact_parts.append(tailored_cv.email)
    if tailored_cv.phone:
        contact_parts.append(tailored_cv.phone)
    if tailored_cv.linkedin_url:
        contact_parts.append(tailored_cv.linkedin_url)
    if tailored_cv.github_url:
        contact_parts.append(tailored_cv.github_url)
    if tailored_cv.portfolio_url:
        contact_parts.append(tailored_cv.portfolio_url)

    if contact_parts:
        sub_header = doc.add_paragraph()
        sub_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_header.paragraph_format.space_after = Pt(8)
        contact_run = sub_header.add_run("  •  ".join(contact_parts))
        contact_run.font.size = Pt(9.5)

    # -------------------------------------------------------------
    # 2. PROFESSIONAL SUMMARY
    # -------------------------------------------------------------
    if tailored_cv.professional_summary:
        _add_section_heading(doc, "Professional Summary")
        summary_p = doc.add_paragraph()
        summary_p.paragraph_format.space_after = Pt(6)
        summary_p.paragraph_format.line_spacing = 1.15
        summary_run = summary_p.add_run(tailored_cv.professional_summary)
        summary_run.font.size = Pt(10)

    # -------------------------------------------------------------
    # 3. SKILLS & COMPETENCIES
    # -------------------------------------------------------------
    if tailored_cv.skills:
        _add_section_heading(doc, "Technical & Professional Skills")
        for skill_group in tailored_cv.skills:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

            cat_run = p.add_run(f"{skill_group.category}: ")
            cat_run.bold = True
            cat_run.font.size = Pt(10)

            skills_run = p.add_run(skill_group.skills)
            skills_run.font.size = Pt(10)

    # -------------------------------------------------------------
    # 4. WORK EXPERIENCE
    # -------------------------------------------------------------
    if tailored_cv.experiences:
        _add_section_heading(doc, "Professional Experience")
        for exp in tailored_cv.experiences:
            # Job Title + Date Line
            p_head = doc.add_paragraph()
            p_head.paragraph_format.space_before = Pt(6)
            p_head.paragraph_format.space_after = Pt(1)
            p_head.paragraph_format.keep_with_next = True

            role_run = p_head.add_run(f"{exp.role} | {exp.company}")
            role_run.bold = True
            role_run.font.size = Pt(10.5)

            if exp.dates:
                # Right-aligned or tabbed dates
                p_head.add_run(f"\t{exp.dates}").italic = True

            if exp.location:
                loc_p = doc.add_paragraph()
                loc_p.paragraph_format.space_after = Pt(2)
                loc_p.paragraph_format.keep_with_next = True
                loc_run = loc_p.add_run(exp.location)
                loc_run.italic = True
                loc_run.font.size = Pt(9.5)

            # Bullet Points
            for bullet in exp.tailored_bullets:
                b_p = doc.add_paragraph(style="List Bullet")
                b_p.paragraph_format.space_after = Pt(2)
                b_p.paragraph_format.line_spacing = 1.15
                b_run = b_p.add_run(bullet)
                b_run.font.size = Pt(10)

    # -------------------------------------------------------------
    # 5. PROJECTS
    # -------------------------------------------------------------
    if tailored_cv.projects:
        _add_section_heading(doc, "Key Projects")
        for proj in tailored_cv.projects:
            p_head = doc.add_paragraph()
            p_head.paragraph_format.space_before = Pt(6)
            p_head.paragraph_format.space_after = Pt(1)
            p_head.paragraph_format.keep_with_next = True

            title_run = p_head.add_run(proj.title)
            title_run.bold = True
            title_run.font.size = Pt(10.5)

            if proj.technologies:
                tech_run = p_head.add_run(f" [{proj.technologies}]")
                tech_run.italic = True
                tech_run.font.size = Pt(9.5)

            for bullet in proj.tailored_bullets:
                b_p = doc.add_paragraph(style="List Bullet")
                b_p.paragraph_format.space_after = Pt(2)
                b_p.paragraph_format.line_spacing = 1.15
                b_run = b_p.add_run(bullet)
                b_run.font.size = Pt(10)

    # -------------------------------------------------------------
    # 6. EDUCATION
    # -------------------------------------------------------------
    if tailored_cv.education:
        _add_section_heading(doc, "Education")
        for edu in tailored_cv.education:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

            deg_run = p.add_run(f"{edu.degree} — {edu.institution}")
            deg_run.bold = True
            deg_run.font.size = Pt(10)

            if edu.graduation_year:
                p.add_run(f" ({edu.graduation_year})").italic = True

    # -------------------------------------------------------------
    # 7. CERTIFICATIONS
    # -------------------------------------------------------------
    if tailored_cv.certifications:
        _add_section_heading(doc, "Certifications")
        for cert in tailored_cv.certifications:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)

            c_run = p.add_run(f"{cert.name} — {cert.issuing_organization}")
            c_run.bold = True
            c_run.font.size = Pt(10)

            if cert.issue_date:
                p.add_run(f" (Issued: {cert.issue_date})")

    # Save to file
    safe_name = "".join(c for c in tailored_cv.full_name if c.isalnum() or c in (" ", "_")).rstrip()
    filename = f"{safe_name.replace(' ', '_')}_Tailored_CV.docx"
    file_path = os.path.join(output_dir, filename)

    doc.save(file_path)
    return file_path